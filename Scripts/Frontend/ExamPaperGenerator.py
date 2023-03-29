from PyQt5.QtWidgets import QMainWindow, QFileDialog
from .Generated.ExamPaperGenerated import Ui_PaperGenerator
from .Util.QuestionPartFunctionHelpers import (GetQuestionAndParts,
                                               GetFullMarkscheme)
from .Util.CriteriaClass import (CriteriaStruct, TOPICS,
                                 COMPONENT1TOPICS, COMPONENT2TOPICS)
from docx import Document
from docx.shared import Pt
from sqlitehandler import SQLiteHandler
from typing import Set, List, Dict
from random import random, randint
from .AlertWindowHandler import AlertWindow
from .OutputWindowHandler import OutputWindowHandler
import os
import subprocess
import shutil
from .ImagesViewHandler import ImagesViewHandler
from .Util.imageClass import AreImagesAvailable
# handles the generation of random questions


class ExamPaperHandler(Ui_PaperGenerator, QMainWindow):

    def __init__(self, parent=None):
        """
        For generating random exam papers.
        """
        super().__init__(parent)
        self.setupUi(self)
        self.SQLSocket = SQLiteHandler()
        # probability of a long answer questions at the end (default 80%)
        self.LONGANSWERPROBABILITY = 0.8
        # margin - if the marks already generated is less than
        # this then add more
        # = 0.4, if min is 80 and max is 120,
        # and current marks is 90, then mid = 100
        # 100 - 90 = 10
        # total range is 40
        # 40 * 0.4 = 16
        # so add another question until current marks is greater
        # than min + 16 = 96
        self.MARKSMARGIN = 0.4
        # min marks for a question to be considered long
        self.MINLONGMARKS = 9
        self.currentQuestionIDs: List[str] = []
        self.TotalMarks: int = 0
        self.SetupInputWidgets()
        self.ConnectSignalSlots()
        self.show()

    def ConnectSignalSlots(self):
        self.pbGenerate.clicked.connect(self.GeneratePaper)
        self.cbLevel.currentTextChanged.connect(self.OnUpdateLevel)
        self.pbOutputToWord.clicked.connect(self.SaveToDoc)
        self.pbOutputTxt.clicked.connect(self.SaveToText)
        self.sbMax.valueChanged.connect(self.ChangeMax)
        self.sbMin.valueChanged.connect(self.ChangeMin)
        self.pbShowMarkscheme.clicked.connect(self.ShowMarkscheme)
        self.pbTextMarkscheme.clicked.connect(self.SaveMarkschemeText)
        self.pbMarkschemeDoc.clicked.connect(self.SaveMarkschemeDoc)
        self.pbShowImages.clicked.connect(self.ShowImages)

    def ShowImages(self):
        """
        Shows images for the current question
        """
        if len(self.currentQuestionIDs) == 0:
            return
        ImagesViewHandler(self.currentQuestionIDs, False, parent=self)

    def OpenFile(self, fileName):
        """
        Try to open the file after saving it.
        """
        # trying a cross-platform method I found on SO
        # https://stackoverflow.com/questions/6178154/open-a-text-file-using-notepad-as-a-help-file-in-python

        if hasattr(os, "startfile"):
            os.startfile(fileName)
        elif shutil.which("xdg-open"):
            subprocess.call(["xdg-open", fileName])
        elif "EDITOR" in os.environ:
            subprocess.call([os.environ["EDITOR"], fileName])

    def SaveMarkschemeDoc(self):
        """
        Saves markscheme to word document
        """
        if len(self.currentQuestionIDs) == 0:
            return
        # save
        try:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getSaveFileName(
                self, "Save File", "", "Word Document(*.docx)",
                options=options)
            if fileName:
                # create document
                document = Document()
                # need to format it properly for document
                # this involves putting each question as its own paragraph.
                style = document.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(12)
                number = 1
                for selectedid in self.currentQuestionIDs:
                    text = GetFullMarkscheme(self.SQLSocket, selectedid,
                                             number)
                    number += 1
                    paragraph = document.add_paragraph(text)
                    paragraph.style = document.styles['Normal']
                    # font and formatting
                    # run = paragraph.add_run()
                    # font = run.font
                    # font.name = "Arial"
                    # font.size = Pt(12)
                paragraph = document.add_paragraph(
                    f"Total marks: {self.TotalMarks}"
                )
                run = paragraph.add_run()
                font = run.font
                font.name = "Arial"
                font.size = Pt(14)
                font.bold = True
                # save the document
                document.save(fileName)
                try:
                    self.OpenFile(fileName)
                except Exception as e:
                    AlertWindow(f"Could not open the saved file: {e}")
        except Exception as e:
            AlertWindow("Could not save: " + str(e))

    def ShowMarkscheme(self):
        """
        Show markscheme window with the paper markscheme
        """
        if len(self.currentQuestionIDs) == 0:
            return
        # get MS text
        mstext = self.GetMarkschemeText()
        output = OutputWindowHandler(
            "Markscheme for generated paper",
            mstext,
            self.currentQuestionIDs,
            parent=self
        )

    def SaveMarkschemeText(self):
        """
        Save markscheme to a text document.
        Not to be confused with GetMarkschemeText
        for getting markscheme text.
        """
        if len(self.currentQuestionIDs) == 0:
            return
        # get text
        mstext = self.GetMarkschemeText()
        try:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getSaveFileName(
                self, "Save File", "", "Text Files(*.txt)",
                options=options)
            if fileName:
                with open(fileName, "w", encoding="utf-8") as f:
                    # save file
                    f.write(mstext)
                try:
                    self.OpenFile(fileName)
                except Exception as e:
                    AlertWindow(f"Saved file could not be opened: {e}")
        except Exception as e:
            AlertWindow("File could not be saved: " + str(e))

    def GetMarkschemeText(self) -> str:
        """
        For generating .txt or just showing text.
        Return None if not available
        """
        if len(self.currentQuestionIDs) == 0:
            return

        text = ""
        number = 1
        for questionID in self.currentQuestionIDs:
            mstext = GetFullMarkscheme(self.SQLSocket, questionID,
                                       number)
            text += mstext
            text += "\n"
            number += 1

        # add total marks at end
        text += f"\nTotal marks: {self.TotalMarks}"
        return text

    def ChangeMin(self):
        """
        Update max so that it cannot be less than min
        """
        self.sbMax.setMinimum(self.sbMin.value())

    def ChangeMax(self):
        """
        Update min so it cannot be greater than max.
        """
        self.sbMin.setMaximum(self.sbMax.value())

    def SaveToText(self):
        """
        Save text edit contents to text file
        """
        if not self.teOutput.toPlainText():
            return
        # it has writing so save
        text = self.teOutput.toPlainText()
        text += f"\nTotal Marks: {self.TotalMarks}"
        try:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getSaveFileName(
                self, "Save File", "", "Text Files(*.txt)",
                options=options)
            if fileName:
                with open(fileName, "w", encoding="utf-8") as f:
                    # save file
                    f.write(text)
                try:
                    self.OpenFile(fileName)
                except Exception as e:
                    AlertWindow(f"Could not open the saved file: {e}")
        except Exception as e:
            AlertWindow(f"Could not save: {e}")

    def SaveToDoc(self):
        """
        Save to docx
        """
        if not self.teOutput.toPlainText():
            return
        # save
        try:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getSaveFileName(
                self, "Save File", "", "Word Document(*.docx)",
                options=options)
            if fileName:
                # create document
                document = Document()
                style = document.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(12)
                # need to format it properly for document
                # this involves putting each question as its own paragraph.
                number = 1
                for selectedid in self.currentQuestionIDs:
                    text = GetQuestionAndParts(self.SQLSocket, selectedid,
                                               number)
                    number += 1
                    paragraph = document.add_paragraph(text)
                    # font and formatting
                    paragraph.style = document.styles['Normal']
                paragraph = document.add_paragraph(
                    f"Total marks: {self.TotalMarks}"
                )
                run = paragraph.add_run()
                font = run.font
                font.name = "Arial"
                font.size = Pt(14)
                font.bold = True
                # save the document
                document.save(fileName)
                try:
                    self.OpenFile(fileName)
                except Exception as e:
                    AlertWindow(f"Could not open the saved file: {e}")
        except Exception as e:
            AlertWindow(f"Could not save file: {e}")

    def SetupInputWidgets(self):
        """
        Setup input widgets as needed
        """
        self.cbLevel.addItem("Both levels")
        self.cbLevel.addItems([
            "A",
            "AS"
        ])
        self.OnUpdateLevel()

    def OnUpdateLevel(self):
        """
        Need to change components available as well
        """
        level = self.cbLevel.currentText()
        components = []
        if level == "A" or self.cbLevel.currentIndex() == 0:
            components.extend([
                "Both components",
                "Component 1",
                "Component 2"
            ])
        elif level == "AS":
            components.append("Component 1")

        self.cbComponent.clear()
        self.cbComponent.addItems(components)

    def GeneratePaper(self):
        """
        Need to generate the exam paper
        We will take a question pool,
        and then randomly generate questions from the pool.
        We can include a long answer question at the end.
        """
        # we want to preferably only have 1 question from each major topic
        # for the component
        # so we will have to overwrite question criteria
        criteria = self.GetQuestionCriteria()
        # we will get all the topics for the component selected
        # using an SQL query.
        availabletopics = []
        if not criteria.component:
            # for if no component is selected
            # componentquery = f"""
            # (Paper.PaperComponent = 'component 1'
            # OR Paper.PaperComponent = 'component 2')
            # """
            # we use a *2 to duplicate all the topics
            # so that there is enough content.
            availabletopics = TOPICS.copy() * 2
        else:
            # componentquery = f"Paper.PaperComponent = '{criteria.component}'"
            if criteria.component == "component 1":
                availabletopics = COMPONENT1TOPICS.copy() * 2
            else:
                availabletopics = COMPONENT2TOPICS.copy() * 2

        levelquery = ""
        if not criteria.level:
            levelquery = f"""
            (Paper.PaperLevel = 'A'
            OR Paper.PaperLevel = 'AS')
            """
        else:
            levelquery = f"Paper.PaperLevel = '{criteria.level}'"

        currentMarks: int = 0
        # store the selected question IDs for use in the paper
        selectedQuestionIDs: List[str] = []
        # if makeing the topic unique then you have to add thi spart first
        # to reserve the topic used by the long answer question.
        # get the long answer question too
        longanswerquery = f"""
        SELECT Question.QuestionID
FROM QUESTION
JOIN QUESTIONTOPIC ON Question.QuestionID = QUESTIONTOPIC.QuestionID
WHERE QuestionTopic.TopicID IN
({','.join([f"'{i}'" for i in availabletopics])})
AND Question.TotalMarks >= {self.MINLONGMARKS}
AND Question.QuestionID IN (SELECT Question.QuestionID FROM
            QUESTION, Parts WHERE NOT EXISTS(
                SELECT Parts.QuestionID FROM Parts
                WHERE Question.QuestionID = Parts.QuestionID)
                )
AND QUESTION.PaperID IN (SELECT Paper.PaperID FROM Paper
WHERE {levelquery.upper()})
    """
        longanswers = self.SQLSocket.queryDatabase(longanswerquery)
        availablelonganswers = list([i[0] for i in longanswers])
        longanswerquestionid = availablelonganswers[
            randint(0, len(availablelonganswers)-1)]
        longanswermarks = f"""
        SELECT TotalMarks FROM Question
        WHERE QuestionID = {longanswerquestionid}
        """
        longanswermarks = self.SQLSocket.queryDatabase(longanswermarks)[
            0][0]
        currentMarks += longanswermarks
        selectedQuestionIDs.append(longanswerquestionid)

        # for a bunch of random topics we get a question then
        # remove it from the available topics
        rangeofmarks = criteria.maxmarks - criteria.minmarks
        # margin calculation
        margin = rangeofmarks * self.MARKSMARGIN
        normmin = criteria.minmarks + margin
        while currentMarks < normmin:
            # add topic
            chosentopic = availabletopics.pop(
                randint(0, len(availabletopics)-1)
            )
            # this query gets all the possible question ids
            availablequestionsidsquery = f"""
            SELECT QuestionID FROM QuestionTopic WHERE
            TopicID = '{chosentopic}' AND QuestionID IN
            (SELECT Question.QuestionID FROM QUESTION
            WHERE QUESTION.PaperID IN (SELECT paper.PaperID FROM
            Paper WHERE {levelquery.upper()}))
            """
            # get the results as a set
            availablequestionids = set(
                [i[0] for i in self.SQLSocket.queryDatabase(
                    availablequestionsidsquery)])
            # remove any questions we have already done
            availablequestionids.difference(selectedQuestionIDs)
            availablequestionids = list(availablequestionids)
            while len(availablequestionids) > 0:
                # keep going until we find a question not
                # exceeding the max marks.
                selectedid = availablequestionids.pop(
                    randint(0, len(availablequestionids)-1)
                )
                # if selected id in the question ids then do
                # not insert
                if (
                    selectedid in selectedQuestionIDs or
                    selectedid in availablelonganswers
                ):
                    continue
                marksquery = f"""
                SELECT TotalMarks FROM Question
                WHERE QuestionID = {selectedid}
                """
                marks = self.SQLSocket.queryDatabase(marksquery)[0][0]
                if currentMarks + marks > criteria.maxmarks:
                    # get next question
                    continue
                else:
                    # question is satisfactory
                    currentMarks += marks
                    # insert at start so long answer is at end
                    selectedQuestionIDs.insert(0, selectedid)
                    break
            # if the topics list is empty
            if len(availabletopics) == 0:
                if currentMarks < criteria.minmarks:
                    # if the min has not been reached
                    # we will output a message to the user
                    alert = AlertWindow("""
                    The paper generated could not exceed
                    the minimum marks. Regenerate the paper to try again.
                    """)
                break

        # we now have all the qeustion ids
        # we can now create the question paper
        self.currentQuestionIDs = selectedQuestionIDs.copy()
        self.TotalMarks = currentMarks
        self.OutputQuestionPaper(selectedQuestionIDs)

    def OutputQuestionPaper(self, selectedQuestionIDs: list):
        """
        Output the question paper to the text edit
        """
        outputtext = ""
        qnumber = 1
        for questionid in selectedQuestionIDs:
            # for each one we get the text.
            outputtext += GetQuestionAndParts(
                self.SQLSocket, questionid, qnumber
            )
            qnumber += 1
            outputtext += "\n"

        # output to textedit
        self.teOutput.setText(outputtext)
        # output marks
        self.lMarksOutput.setText(
            f"Total marks: {self.TotalMarks}")
        # if no images then disable the images button
        # else enable it
        if AreImagesAvailable(selectedQuestionIDs):
            self.pbShowImages.setEnabled(True)
        else:
            self.pbShowImages.setEnabled(False)

    def GetQuestionCriteria(self) -> CriteriaStruct:
        """
        Get the question criteria and return as criteria object.
        Overwriting the parent one as there is the need to change
        the sbMin and sbMax references to default values
        as these boxes are used for the entire paper
        in this class.
        """
        # check component, level, topic to determine
        # if they are on the first option
        # (for all selection)
        # if they are then we replace them with blanks
        # so that we know not to include them as criteria
        # in the SQL query
        component = ""
        if self.cbComponent.currentIndex() != 0:
            component = self.cbComponent.currentText().lower()
        level = ""
        if self.cbLevel.currentIndex() != 0:
            level = self.cbLevel.currentText().lower()

        return CriteriaStruct(
            set(),
            self.sbMin.value(),
            self.sbMax.value(),
            component.lower(),
            level,
            False
        )
